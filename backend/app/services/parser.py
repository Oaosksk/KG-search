import pandas as pd
import json
from pathlib import Path
from docx import Document
from typing import List, Dict, Any

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

class Parser:
    def parse_file(self, file_path: str) -> List[Dict[str, Any]]:
        ext = Path(file_path).suffix.lower()
        
        if ext == '.csv':
            return self._parse_csv(file_path)
        elif ext == '.xlsx':
            return self._parse_excel(file_path)
        elif ext == '.json':
            return self._parse_json(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _parse_csv(self, file_path: str) -> List[Dict[str, Any]]:
        df = pd.read_csv(file_path)
        if df.empty:
            return []
        
        # Convert each row to a content string for better entity extraction
        records = []
        for _, row in df.iterrows():
            # Create a readable content string from the row
            content_parts = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
            if content_parts:
                records.append({
                    'content': ' | '.join(content_parts),
                    **row.to_dict()
                })
        return records if records else df.to_dict('records')
    
    def _parse_excel(self, file_path: str) -> List[Dict[str, Any]]:
        df = pd.read_excel(file_path)
        if df.empty:
            return []
        
        # Convert each row to a content string
        records = []
        for _, row in df.iterrows():
            content_parts = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
            if content_parts:
                records.append({
                    'content': ' | '.join(content_parts),
                    **row.to_dict()
                })
        return records if records else df.to_dict('records')
    
    def _parse_json(self, file_path: str) -> List[Dict[str, Any]]:
        with open(file_path, 'r') as f:
            data = json.load(f)
        return data if isinstance(data, list) else [data]
    
    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Try to split into records if it looks like structured data
        if "Deal ID:" in text or "Client:" in text:
            records = []
            current_record = []
            
            for para in doc.paragraphs:
                line = para.text.strip()
                if not line:
                    if current_record:
                        records.append({"content": "\n".join(current_record)})
                        current_record = []
                else:
                    current_record.append(line)
            
            if current_record:
                records.append({"content": "\n".join(current_record)})
            
            return records if records else [{"content": text}]
        
        return [{"content": text}]
    
    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        if PdfReader is None:
            raise ValueError("PyPDF2 not installed. Install with: pip install PyPDF2")
        reader = PdfReader(file_path)
        text = "\n".join([page.extract_text() for page in reader.pages])
        return [{"content": text}]
    
    def _parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        with open(file_path, 'r') as f:
            return [{"content": f.read()}]

parser = Parser()
